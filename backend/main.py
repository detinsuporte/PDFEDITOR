import os
import io
import shutil
from uuid import uuid4
from typing import List
import io
import json
import zipfile
import pdfplumber
import pandas as pd
import fitz
from PIL import Image
# from pydantic import BaseModel # Not used in the provided OCR code
from pdf2image import convert_from_bytes
from fastapi.concurrency import run_in_threadpool
import docx
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse

from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from pdf2docx import Converter
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from contextlib import asynccontextmanager

MIRIAM_VAULT_DIR = os.path.join(os.path.dirname(__file__), "tmp", "miriam_vault")

def clean_vault():
    if not os.path.exists(MIRIAM_VAULT_DIR):
        return
    count = 0
    for filename in os.listdir(MIRIAM_VAULT_DIR):
        file_path = os.path.join(MIRIAM_VAULT_DIR, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
            count += 1
        except Exception as e:
            print(f"GC failed to delete {file_path}. Reason: {e}")
    print(f"Garbage Collector: {count} arquivos encontrados. Cofre limpo.")

@asynccontextmanager
async def lifespan(app: FastAPI):
    os.makedirs(MIRIAM_VAULT_DIR, exist_ok=True)
    scheduler = AsyncIOScheduler()
    # Runs the security flush every 30 minutes
    scheduler.add_job(clean_vault, 'interval', minutes=30)
    scheduler.start()
    yield
    scheduler.shutdown()

app = FastAPI(title="iLovePDF API Clone", lifespan=lifespan)

# CORS Setup - Permite chamadas do frontend Next.js na porta 3000 e 3001
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "http://127.0.0.1:3000", "http://localhost:3002"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMP_DIR = "temp_uploads"
os.makedirs(TEMP_DIR, exist_ok=True)

def cleanup_files(file_paths: List[str]):
    for path in file_paths:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception as e:
            print(f"Erro ao deletar arquivo temporario: {e}")

@app.post("/api/pdf/add-signature")
async def add_signature_to_pdf(file: UploadFile = File(...), signatures: str = Form(...)):
    """
    Injeta uma assinatura digital (Image Base64) utilizando fitz (PyMuPDF).
    O Front-end envia as coordenadas X, Y, Width e Height normalizadas
    em relação ao Viewport nativo do PDF (1 ponto = 1/72 polegada).
    """
    import base64
    import json
    
    file_bytes = await file.read()
    elements = json.loads(signatures)
    
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        for el in elements:
            page_idx = el['page'] - 1
            if page_idx < 0 or page_idx >= len(doc):
                continue
                
            page = doc[page_idx]
            
            # Extrair os bytes binários da string Base64 "data:image/png;base64,iVBORw0KG..."
            base64_data = el['base64Url']
            if "," in base64_data:
                base64_data = base64_data.split(",")[1]
                
            img_bytes = base64.b64decode(base64_data)
            
            # Bounding Box Vector System
            x0 = el['x']
            y0 = el['y']
            x1 = x0 + el['width']
            y1 = y0 + el['height']
            
            # fitz.Rect() aceita coords left, top, right, bottom nativos.
            rect = fitz.Rect(x0, y0, x1, y1)
            
            # Carimba a assinatura mantendo a transparência do PNG (Alpha overlay automático)
            page.insert_image(rect, stream=img_bytes)

        output_pdf = io.BytesIO()
        doc.save(output_pdf)
        doc.close()
        output_pdf.seek(0)
        
        return StreamingResponse(
            output_pdf,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=Documento_Assinado.pdf"}
        )
    except Exception as e:
        # Zero log policy for passwords dictates no broad exception catching that accidentally logs memory. But `e` is fitz object.
        raise HTTPException(status_code=500, detail="Erro interno no E-Sign")

@app.post("/api/pdf/encrypt")
async def encrypt_pdf(file: UploadFile = File(...), password: str = Form(...)):
    """
    Tranca um documento PDF utilizando criptografia AES de 256 bits nativa do PyMuPDF.
    A senha nunca é impressa ou gravada em disco.
    """
    file_bytes = await file.read()
    
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        output_pdf = io.BytesIO()
        # O AES-256 é automaticamente inferido ao passar permissões de criptografia modernas no PyMuPDF.
        doc.save(
            output_pdf,
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw=password,
            user_pw=password
        )
        doc.close()
        output_pdf.seek(0)
        
        # Overwrite a string de senha na memória para dificultar exploits de despejo
        password = ""
        
        return StreamingResponse(
            output_pdf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=Protegido_{file.filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail="Falha na criptografia do documento.")


@app.post("/api/pdf/decrypt")
async def decrypt_pdf(file: UploadFile = File(...), password: str = Form(...)):
    """
    Recebe um PDF criptografado e uma senha em texto claro.
    Tenta abrir a matriz binária via doc.authenticate(). Se autorizado, exporta o buffer PDF_ENCRYPT_NONE.
    """
    file_bytes = await file.read()
    
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        if not doc.is_encrypted:
            doc.close()
            # Retorna o próprio arquivo se já estava destrancado
            output_pdf = io.BytesIO(file_bytes)
            output_pdf.seek(0)
            return StreamingResponse(output_pdf, media_type="application/pdf")
            
        auth_level = doc.authenticate(password)
        
        if auth_level == 0:
            doc.close()
            password = ""
            raise HTTPException(status_code=401, detail="Senha incorreta ou acesso negado.")
            
        # Re-salva removendo encriptação
        output_pdf = io.BytesIO()
        doc.save(output_pdf, encryption=fitz.PDF_ENCRYPT_NONE)
        doc.close()
        output_pdf.seek(0)
        
        # Purge em memória
        password = ""
        
        return StreamingResponse(
            output_pdf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=Desbloqueado_{file.filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail="Erro interno ao descriptografar documento.")

@app.post("/api/pdf/add-text")
async def add_text_to_pdf_visual(background_tasks: BackgroundTasks, file: UploadFile = File(...), text_elements: str = Form("[]")):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Arquivo deve ser PDF.")
    temp_files = []
    try:
        orig = os.path.join(TEMP_DIR, f"original_{uuid4().hex[:6]}_{file.filename}")
        temp_files.append(orig)
        with open(orig, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        elements = json.loads(text_elements)
        
        doc = fitz.open(orig)
        
        for el in elements:
            page_num = el.get("page", 1) - 1 # 0-indexed in fitz
            if page_num < 0 or page_num >= len(doc):
                continue
            
            page = doc[page_num]
            text = el.get("text", "")
            x = el.get("x", 0)
            y = el.get("y", 0)
            font_size = el.get("font_size", 14)
            font_family = el.get("font_family", "helvetica")
            hex_color = el.get("hex_color", "#000000")
            
            # parse hex_color to rgb (0-1 range)
            hex_color = hex_color.lstrip('#')
            if len(hex_color) == 6:
                r, g, b = tuple(int(hex_color[i:i+2], 16)/255.0 for i in (0, 2, 4))
            else:
                r, g, b = (0, 0, 0)
                
            p = fitz.Point(x, y)
            
            # map standard names
            font_map = {
                "arial": "helv",
                "helvetica": "helv",
                "courier": "cour",
                "times": "ti-ro",
                "times new roman": "ti-ro"
            }
            fitz_font = font_map.get(str(font_family).lower(), "helv")
            
            page.insert_text(p, text, fontsize=font_size, fontname=fitz_font, color=(r,g,b))
            
        out = os.path.join(TEMP_DIR, f"out_{uuid4().hex[:8]}.pdf")
        temp_files.append(out)
        doc.save(out)
        doc.close()
            
        background_tasks.add_task(cleanup_files, temp_files)
        return FileResponse(path=out, filename="Texto_Adicionado.pdf", media_type="application/pdf")
    except Exception as e:
        cleanup_files(temp_files)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/split")
async def split_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...), start_page: int = Form(1), end_page: int = Form(1)):
    if not file.filename.lower().endswith('.pdf'): raise HTTPException(status_code=400)
    temp_files = []
    try:
        orig = os.path.join(TEMP_DIR, f"split_{uuid4().hex[:6]}.pdf")
        temp_files.append(orig)
        with open(orig, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        op = PdfReader(orig)
        writer = PdfWriter()
        total = len(op.pages)
        if start_page < 1 or end_page > total: raise HTTPException(status_code=400, detail="Range invalido")
             
        for i in range(start_page - 1, end_page):
             writer.add_page(op.pages[i])
             
        out = os.path.join(TEMP_DIR, f"out_split_{uuid4().hex[:5]}.pdf")
        temp_files.append(out)
        with open(out, "wb") as ob:
             writer.write(ob)
             
        background_tasks.add_task(cleanup_files, temp_files)
        return FileResponse(path=out, filename="Split.pdf", media_type="application/pdf")
    except Exception as e:
        cleanup_files(temp_files)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/merge")
async def merge_pdfs(files: List[UploadFile] = File(...), order_mapping: str = Form(...)):
    if len(files) < 1: raise HTTPException(status_code=400, detail="Mínimo 1 arquivo")
    try:
        mapping = json.loads(order_mapping)
        doc_dest = fitz.open()
        
        loaded_docs = []
        for file in files:
            file_bytes = await file.read()
            loaded_docs.append(fitz.open(stream=file_bytes, filetype="pdf"))

        for item in mapping:
            f_idx = item.get("fileIndex")
            p_idx = item.get("pageIndex")
            if f_idx is not None and p_idx is not None and 0 <= f_idx < len(loaded_docs):
                doc_src = loaded_docs[f_idx]
                doc_dest.insert_pdf(doc_src, from_page=p_idx, to_page=p_idx)

        for doc in loaded_docs:
            doc.close()

        out_pdf_bytes = doc_dest.write()
        doc_dest.close()

        pdf_stream = io.BytesIO(out_pdf_bytes)
        pdf_stream.seek(0)

        return StreamingResponse(
            pdf_stream,
            media_type="application/pdf",
            headers={
                "Content-Disposition": "attachment; filename=merged.pdf"
            }
        )
    except Exception as e:
        print(f"Merge error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- ENDPOINT DE OCR COM THREADPOOL (EasyOCR) ---

OCR_READER = None

def get_ocr_reader():
    global OCR_READER
    if OCR_READER is None:
        import easyocr
        # Inicia a IA com os idiomas suportados (CPU Mode para contornar GPU/CUDA pesada)
        print("Baixando/Carregando modelos neurais EasyOCR na RAM (Pode demorar 30s na primeira vez)...")
        OCR_READER = easyocr.Reader(['pt', 'en', 'es'], gpu=False)
    return OCR_READER

def perform_ocr_sync(file_path: str, lang: str, filename: str):
    import numpy as np
    """
    Função pesada (CPU-Bound) usando Inteligência Neural EasyOCR (PyTorch).
    Lê diretamente do Cofre Isolado no disco para desafogar a RAM em arquivos gigantes.
    """
    reader = get_ocr_reader()
    
    if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp')):
        # Processamento direto de Imagem nativa do cofre
        img = Image.open(file_path).convert('RGB')
        images = [img]
    else:
        # Processamento de Arquivos PDF pesados do cofre
        with open(file_path, "rb") as f:
            images = convert_from_bytes(f.read(), poppler_path=r"C:\poppler\poppler-24.02.0\Library\bin")
    
    extracted_text = []
    for i, img in enumerate(images):
        # A extração OCR individual
        img_np = np.array(img)
        results = reader.readtext(img_np, detail=0)
        text_page = "\n".join(results)
        extracted_text.append(f"--- PÁGINA {i+1} ---\n\n{text_page}\n")
        
    full_text = "\n".join(extracted_text)
    
    # Criar arquitetura DOCX
    doc = docx.Document()
    doc.add_heading("Extração OCR (Reconhecimento Ótico)", 0)
    
    # Adicionar o texto quebrado corretamente por parágrafos
    for paragraph in full_text.split('\n'):
        doc.add_paragraph(paragraph)
        
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

@app.post("/api/pdf/ocr")
async def extract_ocr(background_tasks: BackgroundTasks, file: UploadFile = File(...), lang: str = Form("por")):
    """
    Descarrega o documento físico no Vault. Despacha OCR. Programa o Garbage Collector 
    via BackgroundTasks para aniquilar o arquivo milissegundos após o Streaming retornar ao cliente.
    """
    file_bytes = await file.read()
    
    vault_path = os.path.join(MIRIAM_VAULT_DIR, f"{uuid4()}_{file.filename}")
    with open(vault_path, "wb") as f:
        f.write(file_bytes)
    
    try:
        doc_io = await run_in_threadpool(perform_ocr_sync, vault_path, lang, file.filename)
        
        # Agenda a auto-destruição (Garbage Collection Imadiato)
        background_tasks.add_task(os.remove, vault_path)
        
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Texto_Extraido_OCR.docx"}
        )
    except Exception as e:
        # Trata a limpeza em caso de falha no Threadpool
        if os.path.exists(vault_path):
            os.remove(vault_path)
            
        print(f"OCR error: {e}")
        # Retorna o erro detalhado via string
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=500, content={"detail": str(e), "error": "ocr_failed"})

@app.post("/api/split/selected")
async def split_selected_pdfs(file: UploadFile = File(...), selected_pages: str = Form(...)):
    try:
        pages_to_keep = json.loads(selected_pages) # [0, 1, 4] -> 0-based
        if not pages_to_keep:
            raise HTTPException(status_code=400, detail="Nenhuma página selecionada.")
            
        file_bytes = await file.read()
        doc_src = fitz.open(stream=file_bytes, filetype="pdf")
        
        doc_dest = fitz.open()

        # Insert only requested pages
        for p in sorted(pages_to_keep):
            if 0 <= p < doc_src.page_count:
                doc_dest.insert_pdf(doc_src, from_page=p, to_page=p)

        doc_src.close()

        out_pdf_bytes = doc_dest.write()
        doc_dest.close()

        pdf_stream = io.BytesIO(out_pdf_bytes)
        pdf_stream.seek(0)

        filename = file.filename or "document"
        new_filename = f"{os.path.splitext(filename)[0]}_extraido.pdf"

        return StreamingResponse(
            pdf_stream,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=\"{new_filename}\""
            }
        )
    except Exception as e:
        print(f"Split selected error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/split/all")
async def split_all_pdfs(file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
        doc_src = fitz.open(stream=file_bytes, filetype="pdf")
        
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            filename_base = os.path.splitext(file.filename or "document")[0]
            
            for p in range(doc_src.page_count):
                temp_doc = fitz.open()
                temp_doc.insert_pdf(doc_src, from_page=p, to_page=p)
                single_page_bytes = temp_doc.write()
                temp_doc.close()
                
                # Zip entry name
                page_name = f"{filename_base}_pag_{p + 1}.pdf"
                zip_file.writestr(page_name, single_page_bytes)

        doc_src.close()
        
        zip_buffer.seek(0)

        zip_filename = f"{filename_base}_dividido.zip"

        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={
                "Content-Disposition": f"attachment; filename=\"{zip_filename}\""
            }
        )
    except Exception as e:
        print(f"Split all error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/to-word")
async def to_word(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.pdf'): raise HTTPException(status_code=400)
    temp_files = []
    try:
        orig = os.path.join(TEMP_DIR, f"word_{uuid4().hex[:6]}.pdf")
        temp_files.append(orig)
        with open(orig, "wb") as b: shutil.copyfileobj(file.file, b)
            
        out = os.path.join(TEMP_DIR, f"word_{uuid4().hex[:6]}.docx")
        temp_files.append(out)
        
        # Converter PDF2DOCX
        cv = Converter(orig)
        cv.convert(out, start=0, end=None)
        cv.close()
        
        background_tasks.add_task(cleanup_files, temp_files)
        return FileResponse(path=out, filename="Convertido.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        cleanup_files(temp_files)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/to-excel")
def to_excel(file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.pdf'): raise HTTPException(status_code=400)
    try:
        # Lẽ o arquivo de forma síncrona dentro da Thread
        file_bytes = file.file.read()
        
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for idx, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for t_idx, table in enumerate(tables):
                        # Se existe cabecalho, usa a primeira linha
                        if len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                        else:
                            df = pd.DataFrame(table)
                            
                        sheet_name = f"Pg_{idx+1}_Tb_{t_idx+1}"
                        df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
            
            # Se nenhuma tabela for encontrada
            if len(writer.sheets) == 0:
                pd.DataFrame([["Nenhuma tabela estruturada foi detectada no PDF."]]).to_excel(writer, sheet_name="Aviso", index=False, header=False)
                
        excel_buffer.seek(0)
        
        filename = file.filename or "document"
        new_filename = f"{os.path.splitext(filename)[0]}_excel.xlsx"

        return StreamingResponse(
            excel_buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename=\"{new_filename}\""
            }
        )
    except Exception as e:
        print(f"To Excel error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/compress")
async def compress_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.pdf'): raise HTTPException(status_code=400)
    try:
        file_bytes = await file.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        # garbage=4 elimina objetos nao usados, otimiza fontes, imagens, etc
        # deflate=True comprime todos os streams
        out_pdf_bytes = doc.write(garbage=4, deflate=True)
        doc.close()

        pdf_stream = io.BytesIO(out_pdf_bytes)
        pdf_stream.seek(0)
        
        filename = file.filename or "document"
        new_filename = f"{os.path.splitext(filename)[0]}_comprimido.pdf"

        return StreamingResponse(
            pdf_stream,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=\"{new_filename}\""
            }
        )
    except Exception as e:
        print(f"Compress error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
def read_root(): return {"status": "success"}

